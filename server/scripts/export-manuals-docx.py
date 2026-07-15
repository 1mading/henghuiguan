#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将恒慧管操作手册导出为 Word（全角色合订本 + 可选分角色版本）。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError:
    print('请先安装: pip install python-docx', file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / 'docs' / '操作手册'
OUT_DIR = DOCS

ROLE_OVERVIEW = """
## 角色说明

| 角色 | 系统标识 | 定位 | 数据范围 |
|------|----------|------|----------|
| **总经理** | gm | 全局统筹、审核决策 | 全公司所有项目与任务 |
| **管理员** | admin | 系统运维、与总经理同等权限 | 全公司所有项目与任务 |
| **部门经理** | manager | 本部门项目落地、任务分配与监督 | 全公司项目只读；本人负责/创建的项目可管理 |
| **执行人员** | staff | 任务执行、进度反馈 | 本人负责/协办任务；项目只读；可被指定为项目团队成员 |

> 总经理与管理员权限相同，区别仅在岗位职责。

## 权限速查矩阵

| 功能 | 总经理/管理员 | 部门经理 | 执行人员 |
|------|:-------------:|:--------:|:--------:|
| 工作台（本人任务） | ✓ | ✓ | ✓ |
| 项目管理（浏览） | 全部可管理 | 全公司只读 + 本人负责项目可管理 | 全公司只读 |
| 新建/编辑项目 | ✓ | ✓ 本人负责项目 | ✗ |
| 指定项目团队成员 | ✓ | ✓ 本人负责项目 | ✗ |
| 作为团队成员编辑任务 | — | — | ✓ 被指定后 |
| 任务中心 | ✓ 本人 + 部门页签 | ✓ 本人 + 部门页签 | ✓ 本人 |
| 跨项目任务依赖 | ✓ 配置 | ✓ 本人可编辑任务 | 仅查看 |
| 任务协办人（告知/辅助） | ✓ | ✓ 本人负责项目 | ✗ |
| 滚动大屏 | ✓ | ✗ | ✗ |
| 团队管理 | ✓ 全员 | ✓ 本部门 | ✗ |
| 人员档案 | ✓ 全员 | ✓ 本部门有限编辑 | ✗ |
| 钉钉通讯录同步 | ✓ | ✗ | ✗ |
| 数据导入 | ✓ | ✓ | ✗ |
| 更新记录 | ✓ | ✓ | ✓ |
| 快速建临时任务 | ✓ | ✓ | ✓ |
"""


def set_doc_style(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_title(doc: Document, text: str, level: int = 0) -> None:
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    else:
        doc.add_heading(text, level=min(level, 3))


def add_paragraph(doc: Document, text: str, *, quote: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    if quote:
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run.italic = True
    elif bold:
        run = p.add_run(text)
        run.bold = True
    else:
        add_rich_text(p, text)


def add_rich_text(paragraph, text: str) -> None:
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ''
            cell = table.rows[i].cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            add_rich_text(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def md_to_docx(doc: Document, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == '---':
            doc.add_paragraph('')
            i += 1
            continue

        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()
            add_title(doc, title, level)
            i += 1
            continue

        if stripped.startswith('>'):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote_lines.append(lines[i].strip().lstrip('>').strip())
                i += 1
            add_paragraph(doc, ' '.join(quote_lines), quote=True)
            continue

        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table_rows(table_lines))
            continue

        if re.match(r'^[-*]\s+', stripped):
            items = []
            while i < len(lines) and re.match(r'^[-*]\s+', lines[i].strip()):
                items.append(re.sub(r'^[-*]\s+', '', lines[i].strip()))
                i += 1
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                add_rich_text(p, item)
            continue

        if re.match(r'^\d+\.\s+', stripped):
            items = []
            while i < len(lines) and re.match(r'^\d+\.\s+', lines[i].strip()):
                items.append(re.sub(r'^\d+\.\s+', '', lines[i].strip()))
                i += 1
            for item in items:
                p = doc.add_paragraph(style='List Number')
                add_rich_text(p, item)
            continue

        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            p = doc.add_paragraph()
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            p.paragraph_format.left_indent = Cm(0.8)
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith('#') or nxt == '---' or nxt.startswith('|') or nxt.startswith('>') or re.match(r'^[-*]\s+', nxt) or re.match(r'^\d+\.\s+', nxt) or nxt.startswith('```'):
                break
            para_lines.append(nxt)
            i += 1
        add_paragraph(doc, ' '.join(para_lines))


def read_md(name: str) -> str:
    path = DOCS / name
    if not path.exists():
        raise FileNotFoundError(path)
    return path.read_text(encoding='utf-8')


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    set_doc_style(doc)
    add_title(doc, title, 0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph('')


def add_part_break(doc: Document, part_title: str, role_hint: str) -> None:
    doc.add_page_break()
    add_title(doc, part_title, 1)
    add_paragraph(doc, role_hint, quote=True)
    doc.add_paragraph('')


def build_combined_doc() -> Document:
    doc = Document()
    add_cover(
        doc,
        '恒慧管 · 操作手册（全角色）',
        '部门内部项目管控系统 · 2026-06 版\n正式环境：钉钉工作台 → 恒慧管\n访问地址：https://henghuiguan.handagroup.com/app',
    )
    md_to_docx(doc, ROLE_OVERVIEW)

    add_part_break(doc, '第一篇 通用说明', '适用：总经理、管理员、部门经理、执行人员 — 登录、任务状态、工时、依赖、通知等')
    md_to_docx(doc, read_md('00-通用说明.md'))

    add_part_break(doc, '第二篇 总经理 / 管理员', '适用角色：gm、admin — 全公司项目管理、钉钉同步、滚动大屏、跨项目依赖配置')
    md_to_docx(doc, read_md('01-总经理与管理员.md'))

    add_part_break(doc, '第三篇 部门经理', '适用角色：manager — 本部门项目与任务、团队饱和度、跨项目依赖（本部门任务）')
    md_to_docx(doc, read_md('02-部门经理.md'))

    add_part_break(doc, '第四篇 执行人员', '适用角色：staff — 本人任务执行、进度反馈、阻塞任务处理')
    md_to_docx(doc, read_md('03-执行人员.md'))

    return doc


def build_gm_doc() -> Document:
    doc = Document()
    add_cover(doc, '恒慧管 · 总经理 / 管理员 操作手册', '访问地址：https://henghuiguan.handagroup.com/app')
    md_to_docx(doc, read_md('00-通用说明.md'))
    doc.add_page_break()
    md_to_docx(doc, read_md('01-总经理与管理员.md'))
    return doc


def build_manager_doc() -> Document:
    doc = Document()
    add_cover(doc, '恒慧管 · 部门经理 操作手册', '访问地址：https://henghuiguan.handagroup.com/app')
    md_to_docx(doc, read_md('00-通用说明.md'))
    doc.add_page_break()
    md_to_docx(doc, read_md('02-部门经理.md'))
    return doc


def build_staff_doc() -> Document:
    doc = Document()
    add_cover(doc, '恒慧管 · 执行人员 操作手册', '访问地址：https://henghuiguan.handagroup.com/app')
    md_to_docx(doc, read_md('00-通用说明.md'))
    doc.add_page_break()
    md_to_docx(doc, read_md('03-执行人员.md'))
    return doc


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    combined_path = OUT_DIR / '恒慧管-操作手册（全角色）.docx'
    gm_path = OUT_DIR / '恒慧管-总经理与管理员操作手册.docx'
    mgr_path = OUT_DIR / '恒慧管-部门经理操作手册.docx'
    staff_path = OUT_DIR / '恒慧管-执行人员操作手册.docx'

    build_combined_doc().save(str(combined_path))
    build_gm_doc().save(str(gm_path))
    build_manager_doc().save(str(mgr_path))
    build_staff_doc().save(str(staff_path))

    print(f'已生成: {combined_path}')
    print(f'已生成: {gm_path}')
    print(f'已生成: {mgr_path}')
    print(f'已生成: {staff_path}')


if __name__ == '__main__':
    main()
